from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

doc = Document()

style = doc.styles['Normal']
font = style.font
font.name = '微軟正黑體'
font.size = Pt(13)
style.element.rPr.rFonts.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}ascii', '微軟正黑體')

for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

def add_title(text, size=22):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor(0x1a, 0x3c, 0x6e)
    return p

def add_heading_text(text, size=15):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor(0x2c, 0x5f, 0x8a)
    return p

def add_body(text, indent=0):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    run = p.add_run(text)
    run.font.size = Pt(13)
    return p

def add_empty():
    doc.add_paragraph()

# ===== Title =====
add_title("劇場演出司儀稿", 26)
add_body("對象：高一、高二學生")
add_body("▬▬▬▬▬▬▬▬▬▬▬▬▬▬▬▬▬▬▬▬▬▬▬▬▬▬▬▬")
add_empty()

# ===== 一、開場白 =====
add_heading_text("一、開場白（燈暗 → 燈亮，司儀上台）", 16)
add_empty()
add_body("（背景音樂微揚，司儀步出舞台中央）")
add_empty()
add_body("各位老師、各位同學，大家好！", indent=1)
add_empty()
add_body("今天，我們要一起看一個故事。", indent=1)
add_body("一個關於「選擇」的故事。", indent=1)
add_body("一個可能就發生在你我身邊的故事。", indent=1)
add_empty()
add_body("在開始之前，我想先問大家一個問題——", indent=1)
add_body("「如果有一天，你發現自己必須做出一個很困難的決定，", indent=1)
add_body("　你會選擇留下，還是離開？」", indent=1)
add_empty()
add_body("（停頓 2-3 秒，讓觀眾思考）", indent=1)
add_empty()
add_body("還好，今天不用你馬上回答。", indent=1)
add_body("因為接下來的演出，會帶我們一起尋找答案。", indent=1)
add_empty()

# ===== 二、活動目的 =====
add_heading_text("二、活動目的介紹", 16)
add_empty()
add_body("今天的活動，是由＿＿＿＿（主辦單位）所策劃的劇場演出。", indent=1)
add_body("希望透過戲劇的方式，讓我們不只是從課本上認識知識，", indent=1)
add_body("更能在故事裡感受、在角色中思考。", indent=1)
add_empty()
add_body("這也是為什麼今天邀請到＿＿＿＿（劇團名稱），", indent=1)
add_body("為我們帶來這齣＿＿＿＿（演出名稱）。", indent=1)
add_empty()
add_body("接下來的＿＿分鐘，", indent=1)
add_body("請放下手機、打開感官，", indent=1)
add_body("讓自己完全走進故事裡。", indent=1)
add_empty()

# ===== 三、介紹劇團 =====
add_heading_text("三、介紹劇團與演出團隊", 16)
add_empty()
add_body("接下來，我要隆重介紹今天的演出團隊——", indent=1)
add_body("＿＿＿＿劇團！", indent=1)
add_empty()
add_body("（請插入劇團簡介：成立背景、主要作品、演出風格等）", indent=1)
add_empty()
add_body("今天參與演出的演員有：", indent=1)
add_body("＿＿＿＿＿ 飾演 ＿＿＿＿＿", indent=1.5)
add_body("＿＿＿＿＿ 飾演 ＿＿＿＿＿", indent=1.5)
add_body("＿＿＿＿＿ 飾演 ＿＿＿＿＿", indent=1.5)
add_body("＿＿＿＿＿ 飾演 ＿＿＿＿＿", indent=1.5)
add_empty()
add_body("導演：＿＿＿＿＿", indent=1)
add_body("編劇：＿＿＿＿＿", indent=1)
add_empty()
add_body("讓我們以最熱烈的掌聲，歡迎＿＿＿＿劇團！", indent=1)
add_empty()
add_body("（全場鼓掌，燈光轉換，演出開始）", indent=1,)
add_empty()

# ===== 四、演出中場提醒（如有中場休息） =====
add_heading_text("四、中場休息提醒（若有）", 16)
add_empty()
add_body("（中場休息時間，司儀上台）", indent=1)
add_empty()
add_body("上半場的演出到這裡告一段落。", indent=1)
add_body("現在是＿＿分鐘的中場休息時間。", indent=1)
add_body("請大家不要走遠，＿＿＿（時間）準時回到座位上。", indent=1)
add_body("下半場，還有更多精彩在等待大家！", indent=1)
add_empty()

# ===== 五、演出結束總結 =====
add_heading_text("五、演出結束總結", 16)
add_empty()
add_body("（全劇結束，演員鞠躬謝幕，待掌聲稍歇後，司儀上台）", indent=1)
add_empty()
add_body("謝謝＿＿＿＿劇團帶來這麼精彩的演出！", indent=1)
add_body("（帶頭鼓掌）", indent=1)
add_empty()
add_body("同學們，故事結束了，但思考才正要開始。", indent=1)
add_empty()
add_body("還記得開場前我問的那個問題嗎？", indent=1)
add_body("「你會留下，還是離開？」", indent=1)
add_empty()
add_body("剛剛在故事裡，＿＿＿＿（主要角色）面臨了＿＿＿＿（關鍵選擇）。", indent=1)
add_body("他的選擇，帶來了什麼結果？", indent=1)
add_body("如果是你，你會怎麼做？", indent=1)
add_empty()
add_body("這就是戲劇的力量——", indent=1)
add_body("它不說教，而是讓你在故事裡，看見自己。", indent=1)
add_empty()

# ===== 六、互動引導（可選） =====
add_heading_text("六、互動引導（選擇性使用）", 16)
add_empty()
add_body("如果你對今天的故事有感受、有想法，", indent=1)
add_body("歡迎在課堂上或回教室後，和同學、老師一起討論。", indent=1)
add_empty()
add_body("我們也準備了一份＿＿＿＿（學習單／回饋單），", indent=1)
add_body("請各班班長於會後至＿＿＿＿（地點）領取。", indent=1)
add_empty()

# ===== 七、結語 =====
add_heading_text("七、結語", 16)
add_empty()
add_body("最後，", indent=1)
add_body("再次感謝＿＿＿＿劇團的精彩演出，", indent=1)
add_body("感謝＿＿＿＿（協助單位），", indent=1)
add_body("也感謝在座每一位同學，", indent=1)
add_body("謝謝你們的專注與掌聲。", indent=1)
add_empty()
add_body("（停頓一下，語氣溫暖）", indent=1)
add_body("希望今天的故事，", indent=1)
add_body("能在你心裡留下一些什麼。", indent=1)
add_body("也許現在還說不清楚，", indent=1)
add_body("但沒關係，", indent=1)
add_body("好的戲劇，就是這樣——", indent=1)
add_body("它會在你想起來的時候，一直陪著你。", indent=1)
add_empty()
add_body("今天的演出到此結束，", indent=1)
add_body("請同學依序離場，謝謝大家！", indent=1)
add_empty()
add_body("（燈亮，散場音樂下）", indent=1)

output = os.path.expanduser(r"~\Documents\emoji-collection\劇場司儀稿.docx")
doc.save(output)
print(f"Word file saved: {output}")
